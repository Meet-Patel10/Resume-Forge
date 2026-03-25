"""
Resume file parser — extracts text from PDF, DOCX, and TXT uploads.
"""
import os


def parse_resume_file(file_storage):
    """Parse an uploaded resume file and return its text content.

    Args:
        file_storage: werkzeug FileStorage object from Flask upload

    Returns:
        dict with 'text' (extracted content) and 'filename'

    Raises:
        ValueError: if file type is not supported
    """
    filename = file_storage.filename or ''
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.pdf':
        return {
            'text': _parse_pdf(file_storage),
            'filename': filename,
            'format': 'pdf',
        }
    elif ext in ('.docx', '.doc'):
        return {
            'text': _parse_docx(file_storage),
            'filename': filename,
            'format': 'docx',
        }
    elif ext == '.txt':
        return {
            'text': file_storage.read().decode('utf-8', errors='replace'),
            'filename': filename,
            'format': 'txt',
        }
    else:
        raise ValueError(
            f'Unsupported file format: {ext}. '
            f'Please upload a .pdf, .docx, or .txt file.'
        )


def _parse_pdf(file_storage):
    """Extract text from a PDF file."""
    from PyPDF2 import PdfReader
    import io

    pdf_bytes = file_storage.read()
    reader = PdfReader(io.BytesIO(pdf_bytes))

    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return '\n'.join(text_parts)


def _parse_docx(file_storage):
    """Extract text from a DOCX file."""
    from docx import Document
    import io

    docx_bytes = file_storage.read()
    doc = Document(io.BytesIO(docx_bytes))

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return '\n'.join(text_parts)
