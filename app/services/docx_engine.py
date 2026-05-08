"""Generate .docx files from structured resume JSON."""

from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def render_docx(resume_data: dict) -> bytes:
    """Turn resume JSON into a .docx file, returns bytes."""
    if not DOCX_AVAILABLE:
        raise ImportError('python-docx is not installed. Run: pip install python-docx')

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    header = resume_data.get('header', {})

    # Name
    name_para = doc.add_heading(header.get('name', ''), level=0)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in name_para.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Contact line
    contact_parts = []
    if header.get('location'):
        contact_parts.append(header['location'])
    if header.get('phone'):
        contact_parts.append(header['phone'])
    if header.get('email'):
        contact_parts.append(header['email'])
    if header.get('linkedin'):
        contact_parts.append(header['linkedin'])
    if header.get('github'):
        contact_parts.append(header['github'])

    if contact_parts:
        contact = doc.add_paragraph(' | '.join(contact_parts))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.style.font.size = Pt(9)

    if header.get('tagline'):
        tagline = doc.add_paragraph(header['tagline'])
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tagline.runs:
            run.italic = True
            run.font.size = Pt(9)

    # Summary
    summary = resume_data.get('summary', '')
    if summary:
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

    # Skills
    skills = resume_data.get('skills', [])
    if skills:
        doc.add_heading('Technical Skills', level=1)
        for group in skills:
            p = doc.add_paragraph()
            run = p.add_run(f"{group.get('category', '')}: ")
            run.bold = True
            p.add_run(', '.join(group.get('items', [])))

    # Experience
    experience = resume_data.get('experience', [])
    if experience:
        doc.add_heading('Experience', level=1)
        for exp in experience:
            p = doc.add_paragraph()
            run = p.add_run(exp.get('title', ''))
            run.bold = True
            p.add_run(f" — {exp.get('company', '')}")
            if exp.get('location'):
                p.add_run(f", {exp['location']}")
            if exp.get('dates'):
                p.add_run(f"  ({exp['dates']})")

            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)

    # Projects
    projects = resume_data.get('projects', [])
    if projects:
        doc.add_heading('Projects', level=1)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get('name', ''))
            run.bold = True
            if proj.get('tech_stack'):
                p.add_run(f" — {proj['tech_stack']}")

            for bullet in proj.get('bullets', []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)

    # Education
    education = resume_data.get('education', [])
    if education:
        doc.add_heading('Education', level=1)
        for edu in education:
            p = doc.add_paragraph()
            run = p.add_run(edu.get('degree', ''))
            run.bold = True
            p.add_run(f" — {edu.get('school', '')}")
            if edu.get('location'):
                p.add_run(f", {edu['location']}")
            if edu.get('dates'):
                p.add_run(f"  ({edu['dates']})")
            if edu.get('details'):
                doc.add_paragraph(edu['details'])

    # Languages etc.
    other = resume_data.get('other', {})
    if other and other.get('languages'):
        doc.add_heading('Languages', level=1)
        doc.add_paragraph(other['languages'])

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
